import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import requests
import base64
import uuid

st.set_page_config(page_title="AI Thesis Assistant", page_icon="🎓", layout="wide")

st.title("🎓 AI Thesis Assistant")
st.subheader("Оформление по ГОСТ + проверка через GigaChat (Сбер)")

# ─── ТВОИ АКТУАЛЬНЫЕ КЛЮЧИ ─────────────────────────────────────────────────────
CLIENT_ID = "019d0ffe-8561-7638-8151-d347f82de15f"
CLIENT_SECRET = "317e5e32-07df-49e1-9a82-7944c5cdd44e"

# ─── Получение токена GigaChat (ИСПРАВЛЕННАЯ ВЕРСИЯ) ─────────────────────────────
def get_gigachat_token():
    # Client Secret уже является готовым base64-ключом — используем его напрямую
    auth_base64 = CLIENT_SECRET

    rq_uid = str(uuid.uuid4())

    headers = {
        "Content-Type": "application/x-www-form-urlencoded",
        "Accept": "application/json",
        "Authorization": f"Basic {auth_base64}",   # ← правильно
        "RqUID": rq_uid
    }

    data = {"scope": "GIGACHAT_API_PERS"}

    try:
        response = requests.post(
            "https://ngw.devices.sberbank.ru:9443/api/v2/oauth",
            headers=headers,
            data=data,
            verify=False,
            timeout=15
        )

        # Дебаг — покажет точный ответ от Сбера
        st.write("🔑 Ответ от OAuth-сервера:", response.text)

        if response.status_code != 200:
            st.error(f"Ошибка получения токена: {response.status_code}\n"
                     f"Ответ сервера:\n{response.text}\n\n"
                     "Рекомендация: Создай **новые** ключи на https://developers.sber.ru")
            st.stop()

        return response.json()["access_token"]
    except Exception as e:
        st.error(f"Ошибка при запросе токена: {str(e)}")
        st.stop()

# ─── Проверка через GigaChat ───────────────────────────────────────────────────
def check_with_gigachat(text, token):
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "X-Request-ID": str(uuid.uuid4()),
        "X-Client-ID": CLIENT_ID
    }

    payload = {
        "model": "GigaChat",
        "messages": [
            {
                "role": "system",
                "content": "Ты строгий российский преподаватель. Проверяй работы по ГОСТ и ФГОС."
            },
            {
                "role": "user",
                "content": f"""
Проанализируй работу и дай отчёт:

1. Структура:
   - Введение (актуальность, цель, задачи, объект, предмет)
   - Основная часть
   - Заключение
   - Список литературы

2. Замечания:
   - Цель и задачи чёткие?
   - Объект и предмет указаны?
   - Есть практическая часть?
   - Список литературы по ГОСТ?

Текст работы:
{text[:12000]}

Отвечай кратко, структурировано, с эмодзи ✅ ⚠️ ❌
"""
            }
        ],
        "temperature": 0.7,
        "max_tokens": 1200,
        "stream": False
    }

    try:
        response = requests.post(
            "https://gigachat.devices.sberbank.ru/api/v1/chat/completions",
            headers=headers,
            json=payload,
            verify=False,
            timeout=60
        )

        st.write("🤖 Ответ от GigaChat:", response.text)  # дебаг

        if response.status_code != 200:
            st.error(f"GigaChat ошибка {response.status_code}\nОтвет:\n{response.text}")
            st.stop()

        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"Ошибка запроса к GigaChat: {str(e)}")
        st.stop()

# ─── Титульный лист ────────────────────────────────────────────────────────────
def add_title_page(doc, institution, student, group, faculty, department, topic, supervisor, year, work_type):
    title_doc = Document()
    section = title_doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    p = title_doc.add_paragraph(institution.upper())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True

    title_doc.add_paragraph()

    if faculty or department:
        p = title_doc.add_paragraph(f"{faculty}\n{department}".strip())
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    for _ in range(5):
        title_doc.add_paragraph()

    p = title_doc.add_paragraph(work_type)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.bold = True

    for _ in range(3):
        title_doc.add_paragraph()

    p = title_doc.add_paragraph("на тему:")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    p = title_doc.add_paragraph(topic.upper())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.bold = True

    for _ in range(6):
        title_doc.add_paragraph()

    p = title_doc.add_paragraph(f"Выполнил(а): {group}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    p = title_doc.add_paragraph(student.upper())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True

    for _ in range(2):
        title_doc.add_paragraph()

    p = title_doc.add_paragraph("Руководитель:")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    p = title_doc.add_paragraph(supervisor)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True

    for _ in range(5):
        title_doc.add_paragraph()

    p = title_doc.add_paragraph(f"ГОРОД — {year}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    for element in reversed(title_doc.element.body):
        doc.element.body.insert(0, element)

# ─── Форматирование по ГОСТ ─────────────────────────────────────────────────────
def format_gost(doc):
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        paragraph.paragraph_format.line_spacing = 1.5

# ─── СПИСКИ ЗАВЕДЕНИЙ ───────────────────────────────────────────────────────────
russia_universities = [
    "МГУ им. М.В. Ломоносова (Москва)", "СПбГУ (Санкт-Петербург)", "НИУ ВШЭ (Москва)", "МФТИ (Московская область)",
    "МГТУ им. Н.Э. Баумана (Москва)", "РАНХиГС при Президенте РФ (Москва)", "РЭУ им. Г.В. Плеханова (Москва)",
    "СПбПУ Петра Великого (Санкт-Петербург)", "УрФУ им. Б.Н. Ельцина (Екатеринбург)", "КФУ (Казань)",
    "НГУ (Новосибирск)", "ТГУ (Томск)", "ЮФУ (Ростов-на-Дону)", "СФУ (Красноярск)", "ДВФУ (Владивосток)",
    "НИЯУ МИФИ (Москва)", "МАИ (Москва)", "РТУ МИРЭА (Москва)", "Финансовый университет при Правительстве РФ (Москва)",
    "Первый МГМУ им. И.М. Сеченова (Москва)", "РУДН (Москва)", "МИСиС (Москва)", "ТПУ (Томск)",
    "Другие / ввести название вуза вручную"
]

hakassia_universities = [
    "Хакасский государственный университет им. Н.Ф. Катанова (Абакан)",
    "Хакасский технический институт — филиал СФУ (Абакан)",
    "Саяно-Шушенский филиал СФУ (Саяногорск)",
    "Другие / ввести вручную"
]

russia_colleges = [
    "Петровский колледж (Санкт-Петербург)", "Колледж №26 Архитектуры, Дизайна и Реинжиниринга (Москва)",
    "Технологический колледж №34 (Москва)", "Московский финансовый колледж", "Колледж РАНХиГС (Москва)",
    "Колледж МГУ (Москва)", "Московский промышленно-экономический колледж РЭУ им. Плеханова",
    "Колледж связи №54 им. П.М. Вострухина (Москва)", "Волгоградский технологический колледж",
    "Другие / ввести название вручную"
]

hakassia_colleges = [
    "Хакасский политехнический колледж (Абакан)", "Колледж ХГУ им. Н.Ф. Катанова (Абакан)",
    "Абаканский строительный техникум", "Абаканский медицинский колледж",
    "Черногорский горно-строительный техникум", "Другие / ввести вручную"
]

russia_schools = [
    "Физтех-лицей им. П.Л. Капицы (Долгопрудный)", "СУНЦ МГУ (Москва)", "Лицей №239 (Санкт-Петербург)",
    "Лицей «Вторая школа» (Москва)", "Лицей НИУ ВШЭ (Москва)", "Школа №179 (Москва)", "Школа «Летово» (Москва)",
    "Лицей №1535 (Москва)", "Школа №57 (Москва)", "Другие / ввести название школы вручную"
]

hakassia_schools = [
    "МБОУ «Лицей имени Н.Г. Булакина» (Абакан)", "Лицей №7 (Саяногорск)",
    "Хакасская национальная гимназия-интернат им. Н.Ф. Катанова (Абакан)",
    "МБОУ «Гимназия» (Абакан)", "МБОУ «Лицей им. А.Г. Баженова» (Черногорск)",
    "Другие / ввести вручную (с городом)"
]

# ─── БОКОВАЯ ПАНЕЛЬ ───────────────────────────────────────────────────────────────
st.sidebar.header("Учебное заведение")

region = st.sidebar.selectbox("Регион", ["Россия (топ)", "Республика Хакасия"])
education_type = st.sidebar.selectbox("Тип заведения", ["ВУЗ", "Колледж/техникум", "Школа"])

if education_type == "ВУЗ":
    univ_list = hakassia_universities if region == "Республика Хакасия" else russia_universities
    raw = st.sidebar.selectbox("Вуз", univ_list)
    institution = st.sidebar.text_input("Или введите вручную", value=raw if "Другие" not in raw else "")
elif education_type == "Колледж/техникум":
    coll_list = hakassia_colleges if region == "Республика Хакасия" else russia_colleges
    raw = st.sidebar.selectbox("Колледж/техникум", coll_list)
    institution = st.sidebar.text_input("Или введите вручную", value=raw if "Другие" not in raw else "")
else:
    school_list = hakassia_schools if region == "Республика Хакасия" else russia_schools
    raw = st.sidebar.selectbox("Школа/лицей/гимназия", school_list)
    institution = st.sidebar.text_input("Или введите вручную", value=raw if "Другие" not in raw else "")

st.sidebar.header("Данные работы")
student = st.sidebar.text_input("ФИО", "Иванов Иван Иванович")
group = st.sidebar.text_input("Группа / класс", "11А")
faculty = st.sidebar.text_input("Факультет / отделение", "")
department = st.sidebar.text_input("Кафедра / специальность", "")
topic = st.sidebar.text_input("Тема", "Исследование...")
supervisor = st.sidebar.text_input("Руководитель", "Петрова А.А.")
year = st.sidebar.text_input("Год", "2026")
work_type = st.sidebar.selectbox("Тип работы", ["ИНДИВИДУАЛЬНЫЙ ПРОЕКТ", "ИССЛЕДОВАТЕЛЬСКАЯ РАБОТА", "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", "БАКАЛАВРСКАЯ РАБОТА", "ДИПЛОМНАЯ РАБОТА (СПО)", "КУРСОВАЯ РАБОТА"])

# ─── ЗАГРУЗКА ФАЙЛА ───────────────────────────────────────────────────────────────
uploaded_file = st.file_uploader("Загрузите .docx-файл", type=["docx"])

if uploaded_file is not None:
    st.success(f"Файл загружен: **{uploaded_file.name}**")

    try:
        doc_bytes = uploaded_file.read()
        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if st.button("🔍 Проверить через GigaChat + оформить по ГОСТ"):
            with st.spinner("Получаем токен и проверяем работу..."):
                token = get_gigachat_token()
                ai_report = check_with_gigachat(full_text, token)

            st.subheader("Отчёт GigaChat по содержанию")
            st.markdown(ai_report)

            with st.spinner("Оформляем документ..."):
                add_title_page(doc, institution, student, group, faculty, department, topic, supervisor, year, work_type)
                format_gost(doc)

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.success("Готово! Титульный лист добавлен, форматирование применено, отчёт получен.")

                st.download_button(
                    label="📥 Скачать готовый .docx",
                    data=bio,
                    file_name=f"проверено_gigachat_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    except Exception as e:
        st.error(f"Ошибка: {str(e)}")

else:
    st.info("Загрузите .docx-файл для проверки и оформления")
